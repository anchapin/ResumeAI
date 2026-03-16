"""ATS Compatibility Analyzer"""

import io
import re
import time
from typing import Optional

from lib.ats.models import (
    ATSCheckResult,
    ATSIssue,
    IssueSeverity,
    IssueType,
)


# Scoring weights for different issue types
ISSUE_SCORES = {
    IssueType.IMAGES: 15,
    IssueType.TABLES: 12,
    IssueType.TEXT_IN_HEADER_FOOTER: 10,
    IssueType.TEXTBOXES: 10,
    IssueType.TWO_COLUMN_LAYOUT: 10,
    IssueType.SPECIAL_CHARACTERS: 5,
    IssueType.COMPLEX_FORMATTING: 5,
    IssueType.CUSTOM_COLUMNS: 8,
    IssueType.FOOTNOTES: 5,
    IssueType.UNSUPPORTED_FONTS: 5,
}

# Deductions by severity
SEVERITY_DEDUCTIONS = {
    IssueSeverity.CRITICAL: 1.5,
    IssueSeverity.WARNING: 1.2,
    IssueSeverity.INFO: 1.0,
}


class ATSAnalyzer:
    """Analyzes resumes for ATS compatibility"""
    
    def __init__(self):
        self._init_detection_patterns()
    
    def _init_detection_patterns(self):
        """Initialize detection patterns for various ATS issues"""
        # Patterns that might indicate special characters
        self.special_char_pattern = re.compile(
            r'[\u2018\u2019\u201c\u201d\u201e\u2022\u2023\u2043\u2219\u25cf\u25cb\u2752\u2762\u2766]'
        )
        
        # Common section headers to detect
        self.section_headers = re.compile(
            r'^(EXPERIENCE|EDUCATION|SKILLS|WORK\s+EXPERIENCE|EMPLOYMENT|PROJECTS|'
            r'CERTIFICATIONS|REFERENCES|LANGUAGES|INTERESTS|HONORS|AWARDS)\s*$',
            re.IGNORECASE | re.MULTILINE
        )
        
        # Two-column detection (multiple short lines in sequence)
        self.two_col_pattern = re.compile(
            r'^[^\n]{0,30}\n[^\n]{0,30}\n[^\n]{0,30}\n[^\n]{0,30}\n',
            re.MULTILINE
        )
    
    def analyze(self, file_content: bytes, filename: str) -> ATSCheckResult:
        """Analyze a resume file for ATS compatibility
        
        Args:
            file_content: Raw file bytes
            filename: Name of the file (for extension detection)
            
        Returns:
            ATSCheckResult with score, issues, and parsed text
        """
        start_time = time.time()
        
        # Determine file type
        file_type = self._get_file_type(filename)
        
        # Extract text based on file type
        try:
            parsed_text, issues = self._extract_text(file_content, file_type)
        except Exception as e:
            return ATSCheckResult(
                file_type=file_type,
                ats_score=0,
                is_parseable=False,
                word_count=0,
                issues=[ATSIssue(
                    issue_type=IssueType.SPECIAL_CHARACTERS,
                    severity=IssueSeverity.CRITICAL,
                    element="file",
                    description=f"Failed to parse file: {str(e)}",
                    fix="Ensure the file is a valid PDF, DOCX, or TXT file"
                )],
                parsed_text="",
                calculation_time_ms=(time.time() - start_time) * 1000
            )
        
        # Detect additional issues from parsed text
        issues.extend(self._detect_text_issues(parsed_text))
        
        # Calculate word count
        word_count = len(parsed_text.split()) if parsed_text else 0
        
        # Calculate score
        ats_score = self._calculate_score(issues)
        
        calculation_time_ms = (time.time() - start_time) * 1000
        
        return ATSCheckResult(
            file_type=file_type,
            ats_score=ats_score,
            is_parseable=True,
            word_count=word_count,
            issues=issues,
            parsed_text=parsed_text,
            calculation_time_ms=calculation_time_ms
        )
    
    def _get_file_type(self, filename: str) -> str:
        """Determine file type from filename"""
        ext = filename.lower().split('.')[-1]
        if ext == 'pdf':
            return 'pdf'
        elif ext in ('docx', 'doc'):
            return 'docx'
        elif ext == 'txt':
            return 'txt'
        return 'unknown'
    
    def _extract_text(self, file_content: bytes, file_type: str) -> tuple[str, list[ATSIssue]]:
        """Extract text from file based on type"""
        issues = []
        
        if file_type == 'pdf':
            return self._extract_from_pdf(file_content, issues)
        elif file_type == 'docx':
            return self._extract_from_docx(file_content, issues)
        elif file_type == 'txt':
            return file_content.decode('utf-8', errors='replace'), issues
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_content: bytes, issues: list[ATSIssue]) -> tuple[str, list[ATSIssue]]:
        """Extract text from PDF"""
        try:
            import pypdf
        except ImportError:
            # Fall back to simple text extraction
            text = file_content.decode('utf-8', errors='replace')
            issues.append(ATSIssue(
                issue_type=IssueType.COMPLEX_FORMATTING,
                severity=IssueSeverity.WARNING,
                element="pdf",
                description="pypdf not installed - limited analysis possible",
                fix="Install pypdf for better PDF text extraction"
            ))
            return text, issues
        
        pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(text)
                
                # Check for images (by checking if page has resources)
                if '/XObject' in page.get('/Resources', {}):
                    issues.append(ATSIssue(
                        issue_type=IssueType.IMAGES,
                        severity=IssueSeverity.CRITICAL,
                        element=f"Page {page_num}",
                        description="Images detected on page - ATS cannot read them",
                        fix="Remove images or convert to plain text format"
                    ))
                
                # Check for tables (by analyzing text layout)
                if self._has_table_structure(text):
                    issues.append(ATSIssue(
                        issue_type=IssueType.TABLES,
                        severity=IssueSeverity.CRITICAL,
                        element=f"Page {page_num}",
                        description="Table structure detected - ATS may not parse correctly",
                        fix="Convert tables to plain text with clear separators"
                    ))
                
                # Check for text in header/footer regions
                lines = text.split('\n')
                if len(lines) >= 3:
                    first_line = lines[0].strip()
                    last_line = lines[-1].strip()
                    if len(first_line) < 50 and len(last_line) < 50:
                        # Likely header/footer text
                        if first_line:
                            issues.append(ATSIssue(
                                issue_type=IssueType.TEXT_IN_HEADER_FOOTER,
                                severity=IssueSeverity.WARNING,
                                element=f"Page {page_num} - Header",
                                description=f"Potential header text: {first_line[:30]}",
                                fix="Move header content to main body"
                            ))
                        if last_line:
                            issues.append(ATSIssue(
                                issue_type=IssueType.TEXT_IN_HEADER_FOOTER,
                                severity=IssueSeverity.WARNING,
                                element=f"Page {page_num} - Footer",
                                description=f"Potential footer text: {last_line[:30]}",
                                fix="Move footer content to main body or remove"
                            ))
        
        return '\n'.join(text_parts), issues
    
    def _has_table_structure(self, text: str) -> bool:
        """Detect if text likely contains table structure"""
        lines = text.split('\n')
        table_indicators = 0
        
        for line in lines:
            # Check for multiple pipes or tabs in a line
            if '|' in line or '\t' in line:
                table_indicators += 1
            # Check for aligned columns (multiple spaces between text)
            if re.search(r' {2,}[A-Za-z]', line):
                table_indicators += 1
        
        # If significant portion of lines have table indicators
        return table_indicators > len(lines) * 0.2 if lines else False
    
    def _extract_from_docx(self, file_content: bytes, issues: list[ATSIssue]) -> tuple[str, list[ATSIssue]]:
        """Extract text from DOCX"""
        try:
            import docx
        except ImportError:
            text = file_content.decode('utf-8', errors='replace')
            issues.append(ATSIssue(
                issue_type=IssueType.COMPLEX_FORMATTING,
                severity=IssueSeverity.WARNING,
                element="docx",
                description="python-docx not installed - limited analysis possible",
                fix="Install python-docx for better DOCX text extraction"
            ))
            return text, issues
        
        doc = docx.Document(io.BytesIO(file_content))
        text_parts = []
        
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # Check for tables in document
        if len(doc.tables) > 0:
            issues.append(ATSIssue(
                issue_type=IssueType.TABLES,
                severity=IssueSeverity.CRITICAL,
                element=f"Document ({len(doc.tables)} tables)",
                description=f"Found {len(doc.tables)} table(s) in document",
                fix="Convert tables to plain text with clear separators"
            ))
        
        # Check for images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                issues.append(ATSIssue(
                    issue_type=IssueType.IMAGES,
                    severity=IssueSeverity.CRITICAL,
                    element="document",
                    description="Image detected in document - ATS cannot read it",
                    fix="Remove images from document"
                ))
                break
        
        return '\n'.join(text_parts), issues
    
    def _detect_text_issues(self, text: str) -> list[ATSIssue]:
        """Detect issues from parsed text"""
        issues = []
        
        # Check for special characters
        special_matches = self.special_char_pattern.findall(text)
        if len(special_matches) > 5:
            issues.append(ATSIssue(
                issue_type=IssueType.SPECIAL_CHARACTERS,
                severity=IssueSeverity.WARNING,
                element="text",
                description=f"Found {len(special_matches)} special/unusual characters",
                fix="Replace special characters with plain text equivalents"
            ))
        
        # Check for two-column layout indicators
        if self.two_col_pattern.search(text):
            issues.append(ATSIssue(
                issue_type=IssueType.TWO_COLUMN_LAYOUT,
                severity=IssueSeverity.CRITICAL,
                element="layout",
                description="Two-column layout detected - ATS may not parse correctly",
                fix="Use single-column layout for better ATS compatibility"
            ))
        
        # Check for textboxes (common in Word)
        if '■' in text or '□' in text or '▪' in text:
            issues.append(ATSIssue(
                issue_type=IssueType.TEXTBOXES,
                severity=IssueSeverity.WARNING,
                element="formatting",
                description="Text box or special characters detected",
                fix="Convert text boxes to regular paragraphs"
            ))
        
        # Check for footnotes
        if re.search(r'\[\d+\]|\(\d+\)', text):
            issues.append(ATSIssue(
                issue_type=IssueType.FOOTNOTES,
                severity=IssueSeverity.WARNING,
                element="references",
                description="Footnotes or citations detected",
                fix="Include footnote content inline or in main text"
            ))
        
        return issues
    
    def _calculate_score(self, issues: list[ATSIssue]) -> int:
        """Calculate ATS compatibility score (0-100)"""
        score = 100
        
        for issue in issues:
            base_deduction = ISSUE_SCORES.get(issue.issue_type, 5)
            severity_multiplier = SEVERITY_DEDUCTIONS.get(issue.severity, 1.0)
            deduction = base_deduction * severity_multiplier
            score -= deduction
        
        return max(0, int(score))
