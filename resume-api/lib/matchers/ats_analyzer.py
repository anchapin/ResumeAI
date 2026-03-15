"""
ATS Compatibility Analyzer.

Analyzes PDF and DOCX resumes for ATS (Applicant Tracking System) compatibility.
Detects issues that may cause parsing failures in systems like Workday, Greenhouse, iCIMS.
"""

import re
from typing import TypedDict, Literal
from pathlib import Path

import pdfplumber
from docx import Document


class ATSIssue(TypedDict):
    """Represents an ATS compatibility issue."""

    type: Literal["CRITICAL", "WARNING", "INFO"]
    element: str
    description: str
    page: int | None
    fix: str


class ATSAnalysisResult(TypedDict):
    """Result of ATS compatibility analysis."""

    file_type: str
    ats_score: int
    issues: list[ATSIssue]
    parsed_text: str
    word_count: int
    is_parseable: bool


class ATSAnalyzer:
    """Analyzes resumes for ATS compatibility."""

    # Critical issues that cause immediate rejection
    CRITICAL_DEDUCTIONS = {
        "images": 25,
        "scanned_pdf": 25,
        "contact_in_header": 25,
        "contact_in_footer": 25,
        "text_boxes": 25,
        "password_protected": 25,
    }

    # Warning issues that reduce score
    WARNING_DEDUCTIONS = {
        "complex_table": 10,
        "multi_column": 10,
        "non_standard_font": 5,
        "icons_special_bullets": 5,
    }

    # Standard fonts that ATS can parse
    STANDARD_FONTS = {
        "arial",
        "calibri",
        "times new roman",
        "helvetica",
        "georgia",
        "verdana",
        "trebuchet ms",
        "courier new",
    }

    # Standard section headings
    STANDARD_SECTIONS = {
        "work experience",
        "experience",
        "employment",
        "education",
        "skills",
        "certifications",
        "projects",
        "summary",
        "objective",
    }

    def analyze_file(self, filepath: str) -> ATSAnalysisResult:
        """Analyze a resume file for ATS compatibility."""
        path = Path(filepath)
        file_type = path.suffix.lower()

        if file_type == ".pdf":
            return self.analyze_pdf(filepath)
        elif file_type in [".docx", ".doc"]:
            return self.analyze_docx(filepath)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def analyze_pdf(self, filepath: str) -> ATSAnalysisResult:
        """Analyze a PDF resume for ATS compatibility."""
        issues: list[ATSIssue] = []
        parsed_text_parts: list[str] = []
        total_word_count = 0
        is_parseable = True

        try:
            with pdfplumber.open(filepath) as pdf:
                # Check if PDF is encrypted
                if pdf.pages and pdf.pages[0].page_obj:
                    try:
                        # Try to access page content
                        _ = pdf.pages[0].chars
                    except Exception:
                        issues.append(
                            ATSIssue(
                                type="CRITICAL",
                                element="password_protected",
                                description="PDF is password-protected or encrypted",
                                page=None,
                                fix="Remove password protection and re-save the PDF",
                            )
                        )
                        is_parseable = False

                for page_num, page in enumerate(pdf.pages, 1):
                    # Check for scanned PDF (no text)
                    if not page.chars or len(page.chars) == 0:
                        issues.append(
                            ATSIssue(
                                type="CRITICAL",
                                element="scanned_pdf",
                                description="PDF appears to be scanned (no extractable text)",
                                page=page_num,
                                fix="Re-create resume as text-based PDF or use OCR",
                            )
                        )
                        is_parseable = False
                        continue

                    # Check for images
                    if page.images:
                        issues.append(
                            ATSIssue(
                                type="CRITICAL",
                                element="images",
                                description=f"Found {len(page.images)} image(s) on page {page_num}",
                                page=page_num,
                                fix="Remove all images and use text descriptions instead",
                            )
                        )

                    # Check for tables
                    tables = page.find_tables()
                    if tables:
                        for table in tables:
                            if len(table.rows) > 2 or (
                                table.columns and len(table.columns) > 2
                            ):
                                issues.append(
                                    ATSIssue(
                                        type="WARNING",
                                        element="complex_table",
                                        description=f"Complex table detected on page {page_num}",
                                        page=page_num,
                                        fix="Convert tables to plain text with labels",
                                    )
                                )

                    # Check header/footer for contact info
                    page_height = page.height
                    header_zone = page.crop((0, 0, page.width, page_height * 0.15))
                    footer_zone = page.crop(
                        (0, page_height * 0.85, page.width, page_height)
                    )

                    header_text = header_zone.extract_text() or ""
                    footer_text = footer_zone.extract_text() or ""

                    if self._contains_contact_info(header_text):
                        issues.append(
                            ATSIssue(
                                type="CRITICAL",
                                element="contact_in_header",
                                description="Contact information in header may be missed by ATS",
                                page=page_num,
                                fix="Move contact info to main body at top of first page",
                            )
                        )

                    if self._contains_contact_info(footer_text):
                        issues.append(
                            ATSIssue(
                                type="CRITICAL",
                                element="contact_in_footer",
                                description="Contact information in footer may be missed by ATS",
                                page=page_num,
                                fix="Move contact info to main body at top of first page",
                            )
                        )

                    # Extract text
                    text = page.extract_text() or ""
                    parsed_text_parts.append(text)
                    total_word_count += len(text.split())

        except Exception as e:
            issues.append(
                ATSIssue(
                    type="CRITICAL",
                    element="parse_error",
                    description=f"Failed to parse PDF: {str(e)}",
                    page=None,
                    fix="Ensure PDF is not corrupted and try again",
                )
            )
            is_parseable = False

        parsed_text = "\n\n".join(parsed_text_parts)
        ats_score = self._calculate_score(issues)

        return ATSAnalysisResult(
            file_type="pdf",
            ats_score=ats_score,
            issues=issues,
            parsed_text=parsed_text,
            word_count=total_word_count,
            is_parseable=is_parseable,
        )

    def analyze_docx(self, filepath: str) -> ATSAnalysisResult:
        """Analyze a DOCX resume for ATS compatibility."""
        issues: list[ATSIssue] = []
        parsed_text_parts: list[str] = []
        total_word_count = 0
        is_parseable = True

        try:
            doc = Document(filepath)

            # Check for images
            inline_shapes = doc.inline_shapes if hasattr(doc, "inline_shapes") else []
            if inline_shapes:
                issues.append(
                    ATSIssue(
                        type="CRITICAL",
                        element="images",
                        description=f"Found {len(inline_shapes)} image(s)",
                        page=None,
                        fix="Remove all images and use text descriptions instead",
                    )
                )

            # Check for tables
            if doc.tables:
                for table in doc.tables:
                    if len(table.rows) > 2 or (
                        table.columns and len(table.columns) > 2
                    ):
                        issues.append(
                            ATSIssue(
                                type="WARNING",
                                element="complex_table",
                                description="Complex table detected",
                                page=None,
                                fix="Convert tables to plain text with labels",
                            )
                        )

            # Check headers/footers for contact info
            for section in doc.sections:
                header_text = section.header.text if section.header else ""
                footer_text = section.footer.text if section.footer else ""

                if self._contains_contact_info(header_text):
                    issues.append(
                        ATSIssue(
                            type="CRITICAL",
                            element="contact_in_header",
                            description="Contact information in header may be missed by ATS",
                            page=None,
                            fix="Move contact info to main body at top of first page",
                        )
                    )
                    break

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    parsed_text_parts.append(para.text)
                    total_word_count += len(para.text.split())

        except Exception as e:
            issues.append(
                ATSIssue(
                    type="CRITICAL",
                    element="parse_error",
                    description=f"Failed to parse DOCX: {str(e)}",
                    page=None,
                    fix="Ensure DOCX is not corrupted and try again",
                )
            )
            is_parseable = False

        parsed_text = "\n\n".join(parsed_text_parts)
        ats_score = self._calculate_score(issues)

        return ATSAnalysisResult(
            file_type="docx",
            ats_score=ats_score,
            issues=issues,
            parsed_text=parsed_text,
            word_count=total_word_count,
            is_parseable=is_parseable,
        )

    def _contains_contact_info(self, text: str) -> bool:
        """Check if text contains contact information."""
        if not text.strip():
            return False

        # Email pattern
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        if re.search(email_pattern, text):
            return True

        # Phone pattern (various formats)
        phone_pattern = r"\b(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}\b"
        if re.search(phone_pattern, text):
            return True

        # LinkedIn URL
        linkedin_pattern = r"linkedin\.com/in/"
        if re.search(linkedin_pattern, text, re.IGNORECASE):
            return True

        return False

    def _calculate_score(self, issues: list[ATSIssue]) -> int:
        """Calculate ATS compatibility score (0-100)."""
        base_score = 100
        deductions = 0

        for issue in issues:
            if issue["type"] == "CRITICAL":
                deductions += self.CRITICAL_DEDUCTIONS.get(issue["element"], 25)
            elif issue["type"] == "WARNING":
                deductions += self.WARNING_DEDUCTIONS.get(issue["element"], 10)
            elif issue["type"] == "INFO":
                deductions += 5

        return max(0, base_score - deductions)


# Singleton instance
_analyzer_instance: ATSAnalyzer | None = None


def get_analyzer() -> ATSAnalyzer:
    """Get singleton ATSAnalyzer instance."""
    global _analyzer_instance
    if _analyzer_instance is None:
        _analyzer_instance = ATSAnalyzer()
    return _analyzer_instance
