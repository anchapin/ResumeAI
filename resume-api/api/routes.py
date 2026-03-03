"""
FastAPI routes for Resume API.
"""

import asyncio
import io
import os
import re
import sys
import zipfile
import hashlib
from pathlib import Path

from fastapi import APIRouter, HTTPException, Request, Response, UploadFile, File, status, WebSocket, Depends
from lib.utils.cache import get_cache_manager, cached
from api.websocket import handle_websocket_connection
from config.dependencies import get_current_user_ws
from database import User
from pydantic import BaseModel, Field
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import httpx  # HTTP client for LinkedIn API
from pypdf import PdfReader  # pypdf for PDF parsing (Python 3.14 compatible)

from .models import (
    ResumeRequest,
    TailorRequest,
    VariantsResponse,
    VariantMetadata,
    TailoredResumeResponse,
    ErrorResponse,
    ResumeData,
)

# Setup library path
lib_path = Path(__file__).parent.parent
sys.path.insert(0, str(lib_path))

from lib.cli import (  # noqa: E402
    ResumeGenerator,
    ResumeTailorer,
    VariantManager,
    CoverLetterGenerator,
)
from lib.utils.validators import (  # noqa: E402
    validate_resume_data,
    validate_file_upload,
    validate_file_content,
)

# Import authentication and rate limiting
from config.dependencies import AuthorizedAPIKey, rate_limit  # noqa: E402
from config import settings  # noqa: E402
from monitoring import logging_config  # noqa: E402

# Configure logging
logger = logging_config.get_logger(__name__)

# Initialize components
LIB_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = LIB_DIR / "templates"

router = APIRouter(prefix="", tags=["Resumes"])

# Initialize managers
_variant_manager = None

def get_variant_manager():
    global _variant_manager
    if _variant_manager is None:
        _variant_manager = VariantManager(str(TEMPLATES_DIR))
    return _variant_manager


@router.get("/health", tags=["Health"])
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "version": "1.0.0"}


@router.post(
    "/render/pdf",
    response_class=Response,
    responses={
        200: {"content": {"application/pdf": {}}, "description": "PDF resume file"},
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Rendering"],
)
@rate_limit(settings.rate_limit_pdf)
async def render_pdf(request: Request, body: ResumeRequest, auth: AuthorizedAPIKey):
    """
    Generate a PDF resume from resume data.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 10 requests per minute per API key.

    Args:
        request: FastAPI Request object
        body: ResumeRequest containing resume_data and variant
        auth: API key authentication info

    Returns:
        PDF file as binary response
    """
    try:
        # Convert Pydantic model to dict
        resume_dict = body.resume_data.model_dump(exclude_none=True)

        # Validate and escape resume data
        resume_dict = validate_resume_data(resume_dict)

        # Validate variant
        if not body.variant or len(body.variant) > 100:
            raise ValueError("Invalid variant name")

        # Initialize generator
        generator = ResumeGenerator(
            templates_dir=str(TEMPLATES_DIR), lib_dir=str(LIB_DIR)
        )

        # Generate PDF
        pdf_bytes = await asyncio.to_thread(
            generator.generate_pdf, resume_data=resume_dict, variant=body.variant
        )

        # Return PDF response
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="resume_{body.variant}.pdf"'
            },
        )

    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF generation failed: {str(e)}",
        )


@router.post(
    "/tailor",
    response_model=TailoredResumeResponse,
    responses={
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Tailoring"],
)
@rate_limit(settings.rate_limit_tailor)
@cached("tailoring:result")
async def tailor_resume(
    request: Request,
    body: TailorRequest,
    auth: AuthorizedAPIKey,
    response: Response,
):
    """
    Tailor a resume to match a job description.
    """
    try:
        # Convert Pydantic model to dict
        resume_dict = body.resume_data.model_dump(exclude_none=True)

        # Validate and escape resume data
        resume_dict = validate_resume_data(resume_dict)

        # Validate tailoring inputs
        if not body.job_description or len(body.job_description) > 50000:
            raise ValueError("Invalid job description")
        if body.company_name and len(body.company_name) > 500:
            raise ValueError("Company name exceeds maximum length")
        if body.job_title and len(body.job_title) > 500:
            raise ValueError("Job title exceeds maximum length")

        # Initialize tailorer
        ai_provider = os.getenv("AI_PROVIDER", "openai")
        tailorer = ResumeTailorer(
            ai_provider=ai_provider,
            api_key=os.getenv(f"{ai_provider.upper()}_API_KEY"),
            model=os.getenv("AI_MODEL"),
        )

        # Tailor resume
        tailored_dict = tailorer.tailor_resume(
            resume_data=resume_dict,
            job_description=body.job_description,
            company_name=body.company_name,
            job_title=body.job_title,
        )

        # Validate tailored data
        tailored_dict = validate_resume_data(tailored_dict)

        # Get keywords
        keywords = tailorer.extract_keywords(body.job_description)

        # Get suggestions
        suggestions = tailorer.suggest_improvements(
            resume_data=resume_dict, job_description=body.job_description
        )

        # Convert back to Pydantic model
        tailored_data = ResumeData(**tailored_dict)

        result = TailoredResumeResponse(
            resume_data=tailored_data, keywords=keywords, suggestions=suggestions
        )

        return result

    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Resume tailoring failed: {str(e)}",
        )


@router.get(
    "/variants",
    response_model=VariantsResponse,
    responses={
        200: {"description": "List of variants"},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Variants"],
)
@rate_limit(settings.rate_limit_variants)
@cached("resume:variants")
async def list_variants(
    request: Request,
    response: Response,
    search: str = None,
    tags: str = None,
    category: str = None,
    industry: str = None,
    layout: str = None,
    color_theme: str = None,
):
    """
    List or filter resume template variants.
    """
    try:
        # Parse tags from comma-separated string
        tags_list = None
        if tags:
            tags_list = [t.strip() for t in tags.split(",") if t.strip()]

        # Use filter if any filter params provided, otherwise get all with metadata
        if any([search, tags_list, category, industry, layout, color_theme]):
            filtered_variants = get_variant_manager().filter_variants(
                search=search,
                tags=tags_list,
                category=category,
                industry=industry,
                layout=layout,
                color_theme=color_theme,
            )
            variant_metadata = [VariantMetadata(**v) for v in filtered_variants]
        else:
            variants = get_variant_manager().list_variants()
            variant_metadata = []
            for variant in variants:
                metadata = get_variant_manager().get_variant_metadata(variant)
                variant_metadata.append(VariantMetadata(**metadata))

        result = VariantsResponse(variants=variant_metadata)
        return result

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list variants: {str(e)}",
        )


@router.get("/", tags=["Root"])
async def root():
    """Root endpoint with API information."""
    return {
        "name": "Resume API",
        "version": "1.0.0",
        "description": "API for generating and tailoring professional resumes",
        "endpoints": {
            "health": "/health",
            "render_pdf": "/render/pdf",
            "tailor": "/tailor",
            "variants": "/variants",
            "docs": "/docs",
        },
    }


# DOCX Export Functions


def create_docx_from_resume(resume_data: dict) -> bytes:
    """
    Generate a DOCX file from resume data.

    Args:
        resume_data: Resume data in JSON Resume format

    Returns:
        DOCX file as bytes
    """
    doc = Document()

    # Get basics
    basics = resume_data.get("basics", {})

    # Title - Name
    if basics.get("name"):
        title = doc.add_heading(basics["name"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Headline
    if basics.get("headline"):
        headline = doc.add_paragraph(basics["headline"])
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info
    contact_parts = []
    if basics.get("email"):
        contact_parts.append(basics["email"])
    if basics.get("phone"):
        contact_parts.append(basics["phone"])
    if basics.get("location"):
        location = basics["location"]
        if isinstance(location, dict):
            location_str = ", ".join(
                filter(None, [location.get("city"), location.get("region")])
            )
        else:
            location_str = str(location)
        if location_str:
            contact_parts.append(location_str)

    if basics.get("url"):
        contact_parts.append(basics["url"])

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Empty line

    # Summary
    if basics.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(basics["summary"])

    # Work Experience
    work = resume_data.get("work", [])
    if work:
        doc.add_heading("Experience", level=1)
        for job in work:
            # Job title and company
            job_title = job.get("position", "")
            company = job.get("company", "")
            if job_title or company:
                job_para = doc.add_paragraph()
                if job_title:
                    run = job_para.add_run(job_title)
                    run.bold = True
                if company:
                    if job_title:
                        job_para.add_run(" at ")
                    run = job_para.add_run(company)
                    run.italic = True

            # Dates
            dates = []
            if job.get("startDate"):
                dates.append(job["startDate"])
            if job.get("endDate"):
                dates.append(job["endDate"])
            elif job.get("current"):
                dates.append("Present")

            if dates:
                date_para = doc.add_paragraph(" - ".join(dates))
                date_para.runs[0].italic = True

            # Description/Summary
            if job.get("summary"):
                doc.add_paragraph(job["summary"])

            # Highlights
            highlights = job.get("highlights", [])
            for highlight in highlights:
                doc.add_paragraph(highlight, style="List Bullet")

            doc.add_paragraph()  # Empty line between jobs

    # Education
    education = resume_data.get("education", [])
    if education:
        doc.add_heading("Education", level=1)
        for edu in education:
            # Degree and institution
            parts = []
            if edu.get("studyType"):
                parts.append(edu["studyType"])
            if edu.get("area"):
                parts.append(edu["area"])
            if edu.get("institution"):
                parts.append(edu["institution"])

            if parts:
                edu_para = doc.add_paragraph()
                run = edu_para.add_run(" - ".join(parts[:2]))
                run.bold = True
                if len(parts) > 2:
                    edu_para.add_run(f" at {parts[2]}")

            # Dates
            dates = []
            if edu.get("startDate"):
                dates.append(edu["startDate"])
            if edu.get("endDate"):
                dates.append(edu["endDate"])

            if dates:
                date_para = doc.add_paragraph(" - ".join(dates))
                date_para.runs[0].italic = True

            doc.add_paragraph()

    # Skills
    skills = resume_data.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=1)

        # Group skills by category if available
        skill_names = []
        for skill in skills:
            if isinstance(skill, dict):
                name = skill.get("name", "")
                if name:
                    skill_names.append(name)
            elif isinstance(skill, str):
                skill_names.append(skill)

        if skill_names:
            doc.add_paragraph(", ".join(skill_names))

    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes.getvalue()


@router.post(
    "/export/docx",
    response_class=Response,
    responses={
        200: {
            "content": {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {}
            },
            "description": "DOCX resume file",
        },
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Export"],
)
@rate_limit(settings.rate_limit_pdf)
async def export_docx(request: Request, body: ResumeRequest, auth: AuthorizedAPIKey):
    """
    Generate a DOCX resume from resume data.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 10 requests per minute per API key.

    Args:
        request: FastAPI Request object
        body: ResumeRequest containing resume_data and variant
        auth: API key authentication info

    Returns:
        DOCX file as binary response
    """
    try:
        # Convert Pydantic model to dict
        resume_dict = body.resume_data.model_dump(exclude_none=True)

        # Generate DOCX
        docx_bytes = await asyncio.to_thread(create_docx_from_resume, resume_dict)

        # Return DOCX response
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="resume_{body.variant or "docx"}.docx"'
            },
        )

    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"DOCX generation failed: {str(e)}",
        )


# PDF Import Endpoint and Helper Functions


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from PDF file.

    Args:
        file_bytes: PDF file content as bytes

    Returns:
        Extracted text content

    Raises:
        ValueError: If PDF is corrupted or invalid
    """
    try:
        # Use pypdf (Python 3.14 compatible alternative to PyMuPDF)
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Invalid or corrupted PDF file: {str(e)}")

    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            text_parts.append(text.strip())

    if not text_parts:
        raise ValueError("No text content found in PDF. This may be a scanned image.")

    return "\n".join(text_parts)


def parse_resume_text(text: str) -> dict:
    """
    Parse extracted text into JSON Resume format.

    Args:
        text: Extracted text from PDF

    Returns:
        Dictionary in JSON Resume format
    """
    lines = text.split("\n")
    resume = {
        "basics": {},
        "work": [],
        "education": [],
        "skills": [],
    }

    # Email pattern
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
    # Phone pattern (various formats)
    phone_pattern = r"(\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"

    # Extract contact info from first few lines (usually name at top)
    name_candidates = []
    for line in lines[:10]:
        line = line.strip()
        if not line:
            continue
        # Look for email
        email_match = re.search(email_pattern, line)
        if email_match and not resume["basics"].get("email"):
            resume["basics"]["email"] = email_match.group()
        # Look for phone
        phone_match = re.search(phone_pattern, line)
        if phone_match and not resume["basics"].get("phone"):
            resume["basics"]["phone"] = phone_match.group()
        # First substantial line is likely the name
        if not resume["basics"].get("name") and len(line) > 2 and len(line) < 50:
            if "@" not in line and not re.search(phone_pattern, line):
                name_candidates.append(line)

    if name_candidates:
        resume["basics"]["name"] = name_candidates[0]

    # Detect sections
    current_section = None

    section_keywords = {
        "experience": [
            "experience",
            "employment",
            "work history",
            "professional experience",
        ],
        "education": ["education", "academic", "degree", "university", "college"],
        "skills": ["skills", "technical skills", "competencies", "technologies"],
        "summary": ["summary", "objective", "profile", "about"],
    }

    work_entries = []
    education_entries = []
    skills_list = []
    summary_text = []

    for line in lines:
        line_lower = line.lower().strip()

        # Check for section headers
        detected_section = None
        for section, keywords in section_keywords.items():
            if any(kw in line_lower for kw in keywords):
                if len(line) < 30:  # Likely a header
                    detected_section = section
                    break

        if detected_section:
            current_section = detected_section
            continue

        # Add content to appropriate section
        if current_section == "experience" and line.strip():
            work_entries.append(line.strip())
        elif current_section == "education" and line.strip():
            education_entries.append(line.strip())
        elif current_section == "skills" and line.strip():
            skills_list.append(line.strip())
        elif current_section == "summary" and line.strip():
            summary_text.append(line.strip())

    # Build work experience entries
    if work_entries:
        # Group consecutive entries into work experiences
        work = []
        current_entry = {}
        for entry in work_entries:
            # Check if this looks like a date line
            if re.match(
                r"\d{4}\s*[-–]\s*\d{4}|\d{4}\s*[-–]\s*present", entry, re.IGNORECASE
            ):
                if current_entry:
                    work.append(current_entry)
                current_entry = {
                    "position": "",
                    "company": "",
                    "startDate": "",
                    "endDate": "",
                    "highlights": [],
                }
                current_entry["startDate"], current_entry["endDate"] = extract_dates(
                    entry
                )
            elif current_entry is not None:
                if not current_entry.get("company"):
                    current_entry["company"] = entry
                elif not current_entry.get("position"):
                    current_entry["position"] = entry
                else:
                    if "highlights" not in current_entry:
                        current_entry["highlights"] = []
                    current_entry["highlights"].append(entry)

        if current_entry and current_entry.get("company"):
            work.append(current_entry)

        if work:
            resume["work"] = work

    # Build education entries
    if education_entries:
        education = []
        for entry in education_entries:
            edu = {}
            # Check for degree
            degree_keywords = ["phd", "master", "bachelor", "associate", "degree"]
            if any(kw in entry.lower() for kw in degree_keywords):
                edu["studyType"] = entry
            else:
                edu["institution"] = entry

            # Check for dates
            dates = extract_dates(entry)
            if dates[0]:
                edu["startDate"] = dates[0]
            if dates[1]:
                edu["endDate"] = dates[1]

            if edu:
                education.append(edu)

        if education:
            resume["education"] = education

    # Build skills section
    if skills_list:
        # Parse skills - could be comma-separated or newline-separated
        all_skills = []
        for skill_line in skills_list:
            # Split by common separators
            parts = re.split(r"[,;|\n]", skill_line)
            for part in parts:
                part = part.strip()
                if part and len(part) < 50:
                    all_skills.append(part)

        if all_skills:
            resume["skills"] = [{"name": s} for s in all_skills[:20]]  # Limit to 20

    # Add summary if found
    if summary_text:
        resume["basics"]["summary"] = " ".join(summary_text[:3])

    return resume


def extract_dates(text: str) -> tuple:
    """
    Extract start and end dates from text.

    Args:
        text: Text containing date information

    Returns:
        Tuple of (start_date, end_date)
    """
    # Pattern for year ranges like "2020-2024" or "2020 - Present"
    date_pattern = r"(\d{4})\s*[-–]\s*(\d{4}|present|current)"
    match = re.search(date_pattern, text, re.IGNORECASE)

    if match:
        start_date = match.group(1)
        end_date = match.group(2)
        if end_date.lower() in ["present", "current"]:
            end_date = ""
        return (start_date, end_date)

    return ("", "")


@router.post(
    "/import/pdf",
    response_model=ResumeData,
    responses={
        200: {"model": ResumeData, "description": "Imported resume data"},
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Import"],
)
@rate_limit("10/minute")
async def import_pdf(
    request: Request, file: UploadFile = File(...), auth: AuthorizedAPIKey = None
):
    """
    Import resume from PDF file.

    Accepts PDF file uploads and extracts resume data in JSON Resume format.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 10 requests per minute per API key.

    Args:
        request: FastAPI Request object
        file: PDF file to import
        auth: API key authentication info (optional)

    Returns:
        ResumeData in JSON Resume format
    """
    try:
        # Read file content
        content = await file.read()

        # Validate file upload
        is_valid, error = validate_file_upload(
            file.filename, len(content), file.content_type
        )
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error,
            )

        # Validate file content
        is_valid, error = validate_file_content(content, ".pdf")
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error,
            )

        # Extract text from PDF
        text = await asyncio.to_thread(extract_text_from_pdf, content)

        # Parse into JSON Resume format
        resume_data = parse_resume_text(text)

        # Validate and escape data
        resume_data = validate_resume_data(resume_data)

        # Validate and return
        return ResumeData(**resume_data)

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF import failed: {str(e)}",
        )


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Invalid or corrupted DOCX file: {str(e)}")

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    if not text_parts:
        raise ValueError("No text content found in DOCX file.")

    return "\n".join(text_parts)


@router.post(
    "/import/docx",
    response_model=ResumeData,
    responses={
        200: {"model": ResumeData, "description": "Imported resume data"},
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Import"],
)
@rate_limit("10/minute")
async def import_docx(
    request: Request, file: UploadFile = File(...), auth: AuthorizedAPIKey = None
):
    """
    Import resume from DOCX file.

    Accepts DOCX file uploads and extracts resume data in JSON Resume format.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 10 requests per minute per API key.
    """
    try:
        # Read file content
        content = await file.read()

        # Validate file upload
        is_valid, error = validate_file_upload(
            file.filename, len(content), file.content_type
        )
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error,
            )

        # Validate file content
        is_valid, error = validate_file_content(content, ".docx")
        if not is_valid:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error,
            )

        # Extract text from DOCX
        text = await asyncio.to_thread(extract_text_from_docx, content)

        # Parse into JSON Resume format
        resume_data = parse_resume_text(text)

        # Validate and escape data
        resume_data = validate_resume_data(resume_data)

        # Validate and return
        return ResumeData(**resume_data)

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"DOCX import failed: {str(e)}",
        )


# Request model for LinkedIn import
class LinkedInImportRequest(BaseModel):
    url: str


async def fetch_linkedin_profile(url: str, api_key: str = None) -> dict:
    """
    Fetch LinkedIn profile data using a third-party API.

    This uses the LinkedIn Profile Scraper API as an example.
    The API key should be configured via LINKEDIN_SCRAPER_API_KEY environment variable.
    """
    api_key = api_key or os.getenv("LINKEDIN_SCRAPER_API_KEY")

    if not api_key:
        raise ValueError(
            "LinkedIn profile import is not configured. "
            "Please set LINKEDIN_SCRAPER_API_KEY environment variable."
        )

    # Example API endpoint - this would need to be configured with actual service
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                "https://api.linkedinprofile.io/v1/scrape",
                json={"url": url},
                headers={"Authorization": f"Bearer {api_key}"},
                timeout=30.0,
            )

            if response.status_code == 429:
                raise ValueError("Rate limited. Please try again later.")
            elif response.status_code == 404:
                raise ValueError("LinkedIn profile not found or is private.")
            elif response.status_code != 200:
                raise ValueError(f"Failed to fetch LinkedIn profile: {response.text}")

            return response.json()

        except httpx.TimeoutException:
            raise ValueError("Request timed out. Please try again.")
        except Exception as e:
            raise ValueError(f"Failed to fetch LinkedIn profile: {str(e)}")


def parse_linkedin_to_resume(profile_data: dict) -> dict:
    """Convert LinkedIn profile data to JSON Resume format."""
    resume = {
        "basics": {},
        "work": [],
        "education": [],
        "skills": [],
    }

    # Extract basics
    if profile_data.get("fullName"):
        resume["basics"]["name"] = profile_data["fullName"]

    if profile_data.get("headline"):
        resume["basics"]["headline"] = profile_data["headline"]

    if profile_data.get("email"):
        resume["basics"]["email"] = profile_data["email"]

    if profile_data.get("phone"):
        resume["basics"]["phone"] = profile_data["phone"]

    if profile_data.get("summary"):
        resume["basics"]["summary"] = profile_data["summary"]

    # Location
    if profile_data.get("location"):
        location = profile_data["location"]
        if isinstance(location, dict):
            resume["basics"]["location"] = {
                "city": location.get("city", ""),
                "region": location.get("region", ""),
                "countryCode": location.get("countryCode", ""),
            }
        else:
            resume["basics"]["location"] = {"address": str(location)}

    # Extract work experience
    for job in profile_data.get("experience", []):
        work_entry = {}

        if job.get("title"):
            work_entry["position"] = job["title"]

        if job.get("companyName"):
            work_entry["company"] = job["companyName"]

        if job.get("startDate"):
            work_entry["startDate"] = job["startDate"]

        if job.get("endDate"):
            work_entry["endDate"] = job["endDate"]
        elif job.get("current"):
            work_entry["endDate"] = ""

        if job.get("description"):
            work_entry["summary"] = job["description"]

        if job.get("highlights"):
            work_entry["highlights"] = job["highlights"]

        if work_entry:
            resume["work"].append(work_entry)

    # Extract education
    for edu in profile_data.get("education", []):
        edu_entry = {}

        if edu.get("schoolName"):
            edu_entry["institution"] = edu["schoolName"]

        if edu.get("degreeName"):
            edu_entry["studyType"] = edu["degreeName"]

        if edu.get("fieldOfStudy"):
            edu_entry["area"] = edu["fieldOfStudy"]

        if edu.get("startDate"):
            edu_entry["startDate"] = edu["startDate"]

        if edu.get("endDate"):
            edu_entry["endDate"] = edu["endDate"]

        if edu_entry:
            resume["education"].append(edu_entry)

    # Extract skills
    for skill in profile_data.get("skills", []):
        if isinstance(skill, str):
            resume["skills"].append({"name": skill})
        elif isinstance(skill, dict) and skill.get("name"):
            resume["skills"].append({"name": skill["name"]})

    return resume


@router.post(
    "/import/linkedin",
    response_model=ResumeData,
    responses={
        200: {"model": ResumeData, "description": "Imported resume data"},
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Import"],
)
@rate_limit("5/minute")
async def import_linkedin(
    request: Request, body: LinkedInImportRequest, auth: AuthorizedAPIKey = None
):
    """
    Import resume from LinkedIn profile.

    Accepts LinkedIn profile URL and fetches profile data using configured API.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 5 requests per minute per API key.

    Configuration:
    - LINKEDIN_SCRAPER_API_KEY: API key for LinkedIn scraping service
    """
    # Validate URL
    linkedin_url = body.url.strip()

    # Check if it's a valid LinkedIn URL
    if "linkedin.com" not in linkedin_url.lower():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid LinkedIn profile URL.",
        )

    # Check for common LinkedIn URL patterns
    if not re.search(r"linkedin\.com/(in|pub|profile)/", linkedin_url):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid LinkedIn profile URL format.",
        )

    try:
        # Fetch profile data
        profile_data = await fetch_linkedin_profile(linkedin_url)

        # Parse to JSON Resume format
        resume_data = parse_linkedin_to_resume(profile_data)

        # Return validated data
        return ResumeData(**resume_data)

    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"LinkedIn import failed: {str(e)}",
        )


# LinkedIn JSON File Import Endpoint


@router.post(
    "/import/linkedin-file",
    response_model=ResumeData,
    responses={
        200: {"model": ResumeData, "description": "Imported resume data"},
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Import"],
)
@rate_limit("10/minute")
async def import_linkedin_file(
    request: Request,
    files: List[UploadFile] = File(...),
    auth: AuthorizedAPIKey = None,
):
    """
    Import resume from LinkedIn exported file (JSON, ZIP, or multiple CSV files).

    Accepts multiple LinkedIn export formats:
    - Single JSON file (LinkedIn JSON export)
    - ZIP folder (LinkedIn CSV data export with multiple files like Profile.csv, Positions.csv, etc.)
    - Multiple CSV files (Folder upload containing Profile.csv, Positions.csv, etc.)

    Users can download their LinkedIn data from Settings > Data privacy > Get a copy of your data.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 10 requests per minute per API key.
    """
    # Import LinkedIn library (moved inside function to avoid circular imports if any)
    from lib.linkedin import LinkedInImporter

    importer = LinkedInImporter()
    linkedin_data = {}

    try:
        # Case 1: Single file (JSON or ZIP) or single CSV
        if len(files) == 1:
            file = files[0]
            content_type = file.content_type or ""
            filename = file.filename or ""
            filename_lower = filename.lower()

            is_json = content_type in [
                "application/json",
                "text/json",
            ] or filename_lower.endswith(".json")
            is_zip = content_type in [
                "application/zip",
                "application/x-zip-compressed",
            ] or filename_lower.endswith(".zip")

            # Read content
            content = await file.read()

            # Check size
            max_size = 10 * 1024 * 1024 if is_zip else 5 * 1024 * 1024
            if len(content) > max_size:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File too large. Maximum size is {max_size // (1024 * 1024)}MB.",
                )

            if is_json:
                try:
                    import json

                    linkedin_data = json.loads(content.decode("utf-8"))
                except json.JSONDecodeError as e:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Invalid JSON format: {str(e)}",
                    )
            elif is_zip:
                linkedin_data = await _parse_linkedin_csv_zip(content, importer)
            else:
                # If it's a single file but not JSON/ZIP, assume it's a CSV from a folder upload
                # We need to process it as a single CSV in a collection
                # Or reject if user didn't upload enough files (e.g. just Profile.csv)
                # Let's try to process it as part of a CSV set
                csv_files = {filename: content}
                linkedin_data = _parse_linkedin_csv_files(csv_files)

        # Case 2: Multiple files (CSV folder upload)
        else:
            csv_files = {}
            total_size = 0
            for file in files:
                content = await file.read()
                total_size += len(content)
                if file.filename:
                    csv_files[file.filename] = content

            if total_size > 20 * 1024 * 1024:  # 20MB total limit
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Total upload size exceeds 20MB limit.",
                )

            linkedin_data = _parse_linkedin_csv_files(csv_files)

        if not linkedin_data:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No valid LinkedIn data found in upload.",
            )

        # Parse with the importer
        resume_data_dict = importer.parse_export(linkedin_data, mode="overwrite")

        # Convert to JSON Resume format
        resume_data = convert_linkedin_to_json_resume(resume_data_dict)

        # Return validated data
        return ResumeData(**resume_data)

    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"LinkedIn file import failed: {str(e)}",
        )


def _parse_linkedin_csv_files(csv_files: dict) -> dict:
    """
    Parse a dictionary of filename -> content bytes as LinkedIn CSVs.
    """
    import csv

    linkedin_data = {}
    csv_mappings = {
        "Profile.csv": "profile",
        "Positions.csv": "positions",
        "Education.csv": "education",
        "Skills.csv": "skills",
        "Certifications.csv": "certifications",
        "Email Addresses.csv": "emails",
        "PhoneNumbers.csv": "phones",
    }

    # Process standard files
    for filename, content_bytes in csv_files.items():
        # Check against mappings (case insensitive match)
        for map_file, data_key in csv_mappings.items():
            if filename.lower().endswith(map_file.lower()):
                try:
                    content = content_bytes.decode("utf-8")
                    rows = list(csv.DictReader(io.StringIO(content)))
                    if rows:
                        cleaned_rows = [
                            {k: v.strip() if v else "" for k, v in row.items()}
                            for row in rows
                        ]
                        linkedin_data[data_key] = cleaned_rows
                except Exception as e:
                    print(f"Warning: Could not parse {filename}: {e}")

    # Process Profile Summary
    for filename, content_bytes in csv_files.items():
        if filename.lower().endswith("profile summary.csv"):
            try:
                content = content_bytes.decode("utf-8")
                rows = list(csv.DictReader(io.StringIO(content)))
                if rows and rows[0].get("Summary"):
                    if "profile" not in linkedin_data:
                        linkedin_data["profile"] = []
                    if linkedin_data["profile"] and isinstance(
                        linkedin_data["profile"], list
                    ):
                        linkedin_data["profile"][0]["Summary"] = rows[0]["Summary"]
            except Exception:
                pass

    return _convert_csv_data_to_importer_format(linkedin_data)


async def _parse_linkedin_csv_zip(zip_content: bytes, importer) -> dict:
    """
    Parse LinkedIn CSV data export ZIP file.

    Extracts CSV files from ZIP and combines them into a single data structure
    that can be parsed by LinkedInImporter.

    Args:
        zip_content: Raw ZIP file content
        importer: LinkedInImporter instance for parsing individual CSV files

    Returns:
        Dictionary with parsed LinkedIn data
    """
    import csv

    linkedin_data = {}

    try:
        # Read ZIP file
        zip_buffer = io.BytesIO(zip_content)
        with zipfile.ZipFile(zip_buffer, "r") as zip_file:
            # Define which CSV files to look for and their corresponding data keys
            csv_mappings = {
                "Profile.csv": "profile",
                "Positions.csv": "positions",
                "Education.csv": "education",
                "Skills.csv": "skills",
                "Certifications.csv": "certifications",
                "Email Addresses.csv": "emails",
                "PhoneNumbers.csv": "phones",
            }

            for csv_filename, data_key in csv_mappings.items():
                # Find the CSV file (may be at root or in a subfolder)
                matching_files = [
                    f
                    for f in zip_file.namelist()
                    if f.lower().endswith(csv_filename.lower())
                ]

                if matching_files:
                    try:
                        csv_file_path = matching_files[0]
                        with zip_file.open(csv_file_path) as f:
                            content = f.read().decode("utf-8")
                            rows = list(csv.DictReader(io.StringIO(content)))

                            if rows:
                                # Clean up empty values in rows
                                cleaned_rows = []
                                for row in rows:
                                    cleaned_row = {
                                        k: v.strip() if v else ""
                                        for k, v in row.items()
                                    }
                                    cleaned_rows.append(cleaned_row)

                                linkedin_data[data_key] = cleaned_rows
                    except Exception as e:
                        # Log but continue - some CSV files may be missing
                        print(f"Warning: Could not parse {csv_filename}: {e}")

            # Also check for Profile Summary.csv if present
            profile_summary_files = [
                f
                for f in zip_file.namelist()
                if f.lower().endswith("profile summary.csv")
            ]
            if profile_summary_files:
                try:
                    with zip_file.open(profile_summary_files[0]) as f:
                        content = f.read().decode("utf-8")
                        rows = list(csv.DictReader(io.StringIO(content)))
                        if rows and rows[0].get("Summary"):
                            # Add summary to profile data if it exists
                            if "profile" not in linkedin_data:
                                linkedin_data["profile"] = []
                            if linkedin_data["profile"] and isinstance(
                                linkedin_data["profile"], list
                            ):
                                linkedin_data["profile"][0]["Summary"] = rows[0][
                                    "Summary"
                                ]
                except Exception:
                    pass  # Profile Summary is optional

        # Convert CSV data to format expected by LinkedInImporter
        converted_data = _convert_csv_data_to_importer_format(linkedin_data)
        return converted_data

    except zipfile.BadZipFile as e:
        raise ValueError(f"Invalid ZIP file: {str(e)}")
    except Exception as e:
        raise ValueError(f"Error parsing LinkedIn CSV ZIP: {str(e)}")


def _convert_csv_data_to_importer_format(csv_data: dict) -> dict:
    """
    Convert CSV data structure to format expected by LinkedInImporter.

    CSV data comes as dict of lists (one list per CSV file).
    LinkedInImporter expects fields like 'positions', 'education', 'skills' etc.

    Args:
        csv_data: Dictionary with keys like 'profile', 'positions', 'education'

    Returns:
        Dictionary compatible with LinkedInImporter.parse_export()
    """
    result = {}

    # Handle profile data (usually one row with personal info)
    if "profile" in csv_data and csv_data["profile"]:
        profile_row = csv_data["profile"][0]
        # Map Profile.csv columns to importer field names
        result.update(
            {
                "firstName": profile_row.get("First Name", ""),
                "lastName": profile_row.get("Last Name", ""),
                "headline": profile_row.get("Headline", ""),
                "summary": profile_row.get("Summary", ""),
                "industry": profile_row.get("Industry", ""),
                "locationName": profile_row.get("Geo Location", ""),
            }
        )

        # Handle email addresses (from Email Addresses.csv)
        if csv_data.get("emails"):
            email_list = []
            for email_row in csv_data["emails"]:
                email = email_row.get("Email Address") or email_row.get("email")
                if email:
                    # Prefer confirmed and primary emails
                    is_primary = email_row.get("Primary", "").lower() == "yes"
                    email_list.append({"email": email, "primary": is_primary})

            if email_list:
                # Sort to put primary email first
                email_list.sort(key=lambda x: x["primary"], reverse=True)
                result["emailAddress"] = email_list[0]["email"]

        # Handle phone numbers (from PhoneNumbers.csv)
        if csv_data.get("phones"):
            phone_list = []
            for phone_row in csv_data["phones"]:
                number = phone_row.get("Number") or phone_row.get("number")
                if number:
                    phone_list.append({"phoneNumber": number})
            if phone_list:
                result["phoneNumbers"] = phone_list

    # Handle positions (work experience)
    if "positions" in csv_data:
        positions = []
        for pos in csv_data["positions"]:
            position = {
                "companyName": pos.get("Company Name", ""),
                "title": pos.get("Title", ""),
                "description": pos.get("Description", ""),
                "location": pos.get("Location", ""),
                "timePeriod": {
                    "startDate": _parse_linkedin_date(pos.get("Started On", "")),
                    "endDate": _parse_linkedin_date(pos.get("Finished On", "")),
                },
            }
            if position["companyName"]:  # Only add if company name exists
                positions.append(position)
        if positions:
            result["positions"] = positions

    # Handle education
    if "education" in csv_data:
        education = []
        for edu in csv_data["education"]:
            # Try multiple field names for field of study
            field_of_study = (
                edu.get("Field of Study")
                or edu.get("Activities")
                or edu.get("Notes")
                or ""
            )

            edu_entry = {
                "schoolName": edu.get("School Name", ""),
                "degreeName": edu.get("Degree Name", ""),
                "fieldOfStudy": field_of_study,
                "timePeriod": {
                    "startDate": {"year": _extract_year(edu.get("Start Date", ""))},
                    "endDate": {"year": _extract_year(edu.get("End Date", ""))},
                },
            }
            if edu_entry["schoolName"]:  # Only add if school name exists
                education.append(edu_entry)
        if education:
            result["educations"] = education

    # Handle certifications
    if "certifications" in csv_data:
        certifications = []
        for cert in csv_data["certifications"]:
            cert_entry = {
                "name": cert.get("Name", ""),
                "authority": cert.get("Authority", "") or cert.get("Issuer", ""),
                "timePeriod": {
                    "startDate": _parse_linkedin_date(cert.get("Started On", "")),
                    "endDate": _parse_linkedin_date(cert.get("Finished On", "")),
                },
            }
            if cert_entry["name"]:  # Only add if name exists
                certifications.append(cert_entry)
        if certifications:
            result["certifications"] = certifications

    # Handle skills
    if "skills" in csv_data:
        skills = []
        for skill_row in csv_data["skills"]:
            # Try both "Name" and "Skill" as column names
            skill_name = (
                skill_row.get("Name")
                or skill_row.get("Skill")
                or skill_row.get("skill")
            )
            if skill_name:
                skills.append({"name": skill_name})
        if skills:
            result["skills"] = skills

    return result


def _parse_linkedin_date(date_str: str) -> dict:
    """
    Parse LinkedIn date string to month/year dict format.

    Handles formats like "Apr 2022", "2022-04", "04/2022", "4/2022".

    Args:
        date_str: Date string from LinkedIn CSV

    Returns:
        Dict with 'month' and 'year' keys (compatible with LinkedInImporter format)
    """
    if not date_str or date_str.lower() == "present":
        return {}

    from datetime import datetime

    date_str = date_str.strip()

    try:
        # Try parsing common date formats
        for fmt in ["%b %Y", "%B %Y", "%m/%d/%Y", "%m/%Y", "%Y-%m"]:
            try:
                dt = datetime.strptime(date_str, fmt)
                return {"month": dt.month, "year": dt.year}
            except ValueError:
                continue
    except Exception:
        pass

    return {}


def _extract_year(date_str: str) -> int:
    """Extract year from date string."""
    if not date_str:
        return None

    import re

    # Try to find a 4-digit year
    year_match = re.search(r"(\d{4})", date_str)
    if year_match:
        return int(year_match.group(1))

    return None


def convert_linkedin_to_json_resume(linkedin_data: dict) -> dict:
    """
    Convert parsed LinkedIn data to JSON Resume format.

    Args:
        linkedin_data: Parsed LinkedIn data from LinkedInImporter

    Returns:
        Dictionary in JSON Resume format
    """
    resume = {
        "basics": {},
        "work": [],
        "education": [],
        "skills": [],
        "languages": [],
        "projects": [],
        "certificates": [],
    }

    # Basic info
    if linkedin_data.get("name"):
        resume["basics"]["name"] = linkedin_data["name"]

    if linkedin_data.get("headline") or linkedin_data.get("role"):
        resume["basics"]["label"] = linkedin_data.get("headline") or linkedin_data.get(
            "role"
        )

    if linkedin_data.get("summary"):
        resume["basics"]["summary"] = linkedin_data["summary"]

    if linkedin_data.get("email"):
        resume["basics"]["email"] = linkedin_data["email"]

    if linkedin_data.get("phone"):
        resume["basics"]["phone"] = linkedin_data["phone"]

    if linkedin_data.get("location"):
        resume["location"] = {"city": linkedin_data["location"]}

    # Work experience
    for exp in linkedin_data.get("experience", []):
        work_entry = {
            "company": exp.get("company", ""),
            "position": exp.get("role", ""),
            "startDate": exp.get("startDate", ""),
            "endDate": exp.get("endDate", "") if not exp.get("current") else "",
            "summary": exp.get("description", ""),
        }
        if work_entry["company"] or work_entry["position"]:
            resume["work"].append(work_entry)

    # Education
    for edu in linkedin_data.get("education", []):
        edu_entry = {
            "institution": edu.get("institution", ""),
            "studyType": edu.get("studyType", ""),
            "area": edu.get("area", ""),
            "startDate": edu.get("startDate", ""),
            "endDate": edu.get("endDate", ""),
        }
        if edu_entry["institution"]:
            resume["education"].append(edu_entry)

    # Skills
    for skill in linkedin_data.get("skills", []):
        if isinstance(skill, str):
            resume["skills"].append({"name": skill})
        elif isinstance(skill, dict) and skill.get("name"):
            resume["skills"].append({"name": skill["name"]})

    # Languages
    for lang in linkedin_data.get("languages", []):
        if isinstance(lang, dict):
            resume["languages"].append(
                {
                    "name": lang.get("name", ""),
                    "proficiency": lang.get("proficiency", ""),
                }
            )

    # Projects
    for proj in linkedin_data.get("projects", []):
        if isinstance(proj, dict):
            project_entry = {
                "name": proj.get("name", ""),
                "description": proj.get("description", ""),
                "url": proj.get("url", ""),
            }
            if project_entry["name"]:
                resume["projects"].append(project_entry)

    return resume


# Cover Letter Generation


class CoverLetterRequest(BaseModel):
    """Request to generate a cover letter."""

    resume_data: ResumeData = Field(..., description="Resume data")
    job_description: str = Field(
        ..., min_length=10, max_length=50000, description="Job description text"
    )
    company_name: str = Field(..., max_length=200, description="Company name")
    job_title: str = Field(..., max_length=200, description="Job title")
    tone: str = Field(
        default="professional",
        description="Tone of the cover letter (professional, casual, formal)",
    )


class CoverLetterResponse(BaseModel):
    """Response with generated cover letter."""

    header: str
    introduction: str
    body: str
    closing: str
    full_text: str
    metadata: dict


@router.post(
    "/cover-letter",
    response_model=CoverLetterResponse,
    responses={
        400: {"model": ErrorResponse},
        401: {"model": ErrorResponse},
        403: {"model": ErrorResponse},
        429: {"model": ErrorResponse},
        500: {"model": ErrorResponse},
    },
    tags=["Cover Letter"],
)
@rate_limit("10/minute")
async def generate_cover_letter(
    request: Request, body: CoverLetterRequest, auth: AuthorizedAPIKey = None
):
    """
    Generate a cover letter based on resume and job description.

    Requires API key authentication via X-API-KEY header.

    Rate limit: 10 requests per minute per API key.

    Args:
        request: FastAPI Request object
        body: CoverLetterRequest containing resume_data and job details
        auth: API key authentication info

    Returns:
        CoverLetterResponse with generated cover letter
    """
    try:
        # Convert Pydantic model to dict
        resume_dict = body.resume_data.model_dump(exclude_none=True)

        # Initialize cover letter generator
        ai_provider = os.getenv("AI_PROVIDER", "openai")
        cover_letter_gen = CoverLetterGenerator(
            ai_provider=ai_provider,
            api_key=os.getenv(f"{ai_provider.upper()}_API_KEY"),
            model=os.getenv("AI_MODEL"),
        )

        # Generate cover letter
        cover_letter = cover_letter_gen.generate_cover_letter(
            resume_data=resume_dict,
            job_description=body.job_description,
            company_name=body.company_name,
            job_title=body.job_title,
            tone=body.tone,
        )

        return CoverLetterResponse(**cover_letter)

    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Cover letter generation failed: {str(e)}",
        )


# Test websocket
@router.websocket("/test_ws")
async def test_ws(websocket: WebSocket):
    await websocket.accept()
    await websocket.send_text("hello")
    await websocket.close()


# WebSocket endpoint for real-time collaboration
@router.websocket("/ws/resumes/{resume_id}")
async def websocket_resume(
    websocket: WebSocket,
    resume_id: str,
    current_user_info: tuple[User, float] = Depends(get_current_user_ws),
):
    """
    WebSocket endpoint for real-time collaboration on resumes.
    """
    user, expires_at = current_user_info
    await handle_websocket_connection(
        websocket, resume_id, str(user.id), expires_at=expires_at
    )
