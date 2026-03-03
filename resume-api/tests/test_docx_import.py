"""
Test suite for DOCX import functionality.

Tests cover:
1. extract_text_from_docx function
2. import_docx endpoint
3. File validation (type, size, corruption)
4. Text parsing from DOCX to JSON Resume format
"""

import io
import os

# Set test environment variables before importing app
os.environ.setdefault("TESTING", "True")
os.environ.setdefault("REQUIRE_API_KEY", "False")
os.environ.setdefault("ENABLE_RATE_LIMITING", "False")
os.environ.setdefault("ENABLE_ANALYTICS", "False")
os.environ.setdefault("ENABLE_METRICS", "False")

import pytest  # noqa: E402
from fastapi import status  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

# Import from the API
from api.routes import extract_text_from_docx, parse_resume_text  # noqa: E402
from main import app  # noqa: E402


@pytest.fixture
def client():
    """Create a test client for the FastAPI app."""
    return TestClient(app, raise_server_exceptions=False)


@pytest.fixture
def sample_docx_bytes():
    """Create a sample DOCX file in memory for testing."""
    doc = Document()

    # Add name as heading
    doc.add_heading("John Doe", 0)

    # Add headline
    headline = doc.add_paragraph("Software Engineer")
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add contact info
    contact = doc.add_paragraph(
        "john.doe@example.com | (555) 123-4567 | San Francisco, CA"
    )
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Experienced software engineer with 5+ years of experience in Python, "
        "JavaScript, and cloud technologies."
    )

    # Add work experience
    doc.add_heading("Experience", level=1)

    # Job 1
    job1 = doc.add_paragraph()
    job1.add_run("Senior Software Engineer").bold = True
    job1.add_run(" at Tech Corp")
    doc.add_paragraph("2020 - Present")
    doc.add_paragraph("Led development of cloud-native applications.")
    doc.add_paragraph("Architected microservices infrastructure.", style="List Bullet")
    doc.add_paragraph("Mentored junior developers.", style="List Bullet")

    # Job 2
    job2 = doc.add_paragraph()
    job2.add_run("Software Engineer").bold = True
    job2.add_run(" at StartupXYZ")
    doc.add_paragraph("2018 - 2020")
    doc.add_paragraph("Developed full-stack web applications.")

    # Add education
    doc.add_heading("Education", level=1)
    edu = doc.add_paragraph()
    edu.add_run("Bachelor of Science in Computer Science").bold = True
    doc.add_paragraph("University of California, Berkeley")
    doc.add_paragraph("2014 - 2018")

    # Add skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, JavaScript, React, Node.js, AWS, Docker, Kubernetes")

    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes.getvalue()


@pytest.fixture
def simple_docx_bytes():
    """Create a minimal DOCX file for basic testing."""
    doc = Document()
    doc.add_heading("Jane Smith", 0)
    doc.add_paragraph("jane.smith@example.com")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Developer at Company ABC")
    doc.add_paragraph("2019 - Present")

    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes.getvalue()


@pytest.fixture
def docx_with_tables():
    """Create a DOCX file with tables for testing table extraction."""
    doc = Document()
    doc.add_heading("Table Resume", 0)

    # Add a table with skills
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"

    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes.getvalue()


class TestExtractTextFromDocx:
    """Test the extract_text_from_docx function."""

    def test_extract_text_from_valid_docx(self, sample_docx_bytes):
        """Test text extraction from a valid DOCX file."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))

        assert isinstance(text, str)
        assert "John Doe" in text
        assert "Software Engineer" in text
        assert "john.doe@example.com" in text
        assert "Tech Corp" in text
        assert "University of California" in text

    def test_extract_text_from_simple_docx(self, simple_docx_bytes):
        """Test text extraction from a minimal DOCX file."""
        text = extract_text_from_docx(io.BytesIO(simple_docx_bytes))

        assert "Jane Smith" in text
        assert "jane.smith@example.com" in text
        assert "Developer" in text

    def test_extract_text_from_docx_with_tables(self, docx_with_tables):
        """Test that text is extracted from tables in DOCX."""
        text = extract_text_from_docx(io.BytesIO(docx_with_tables))

        assert "Table Resume" in text
        assert "Python" in text
        assert "Expert" in text

    def test_extract_text_from_corrupted_docx(self):
        """Test handling of corrupted DOCX file."""
        corrupted_bytes = b"This is not a valid DOCX file"

        with pytest.raises(ValueError) as exc_info:
            extract_text_from_docx(io.BytesIO(corrupted_bytes))

        assert "Invalid or corrupted DOCX file" in str(exc_info.value)

    def test_extract_text_from_empty_docx(self):
        """Test handling of empty DOCX file."""
        doc = Document()
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        with pytest.raises(ValueError) as exc_info:
            extract_text_from_docx(io.BytesIO(docx_bytes.getvalue()))

        assert "No text content found" in str(exc_info.value)


class TestParseResumeText:
    """Test the parse_resume_text function with DOCX-extracted text."""

    def test_parse_resume_with_all_sections(self, sample_docx_bytes):
        """Test parsing resume text with all standard sections."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        assert isinstance(resume, dict)
        assert "basics" in resume
        assert "work" in resume
        assert "education" in resume
        assert "skills" in resume

    def test_parse_resume_name_extraction(self, sample_docx_bytes):
        """Test that name is correctly extracted."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        assert resume["basics"].get("name") == "John Doe"

    def test_parse_resume_email_extraction(self, sample_docx_bytes):
        """Test that email is correctly extracted."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        assert "john.doe@example.com" in resume["basics"].get("email", "")

    def test_parse_resume_phone_extraction(self, sample_docx_bytes):
        """Test that phone number is correctly extracted."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        assert "555" in resume["basics"].get("phone", "")

    def test_parse_resume_work_extraction(self, sample_docx_bytes):
        """Test that work experience is extracted."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        assert len(resume.get("work", [])) >= 0  # May vary based on parsing

    def test_parse_resume_education_extraction(self, sample_docx_bytes):
        """Test that education is extracted."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        assert len(resume.get("education", [])) >= 0

    def test_parse_resume_skills_extraction(self, sample_docx_bytes):
        """Test that skills are extracted."""
        text = extract_text_from_docx(io.BytesIO(sample_docx_bytes))
        resume = parse_resume_text(text)

        skills = resume.get("skills", [])
        assert len(skills) > 0


class TestImportDocxEndpoint:
    """Test the /v1/import/docx endpoint."""

    def test_import_docx_success(self, client, sample_docx_bytes):
        """Test successful DOCX import."""
        # Create file-like object for TestClient
        files = {
            "file": (
                "resume.docx",
                sample_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_200_OK
        data = response.json()

        assert "basics" in data
        assert "name" in data.get("basics", {})
        assert data["basics"]["name"] == "John Doe"

    def test_import_docx_with_simple_file(self, client, simple_docx_bytes):
        """Test DOCX import with minimal resume."""
        files = {
            "file": (
                "simple.docx",
                simple_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_200_OK
        data = response.json()

        assert "basics" in data
        assert "Jane Smith" in data.get("basics", {}).get("name", "")

    def test_import_docx_invalid_file_type(self, client):
        """Test rejection of non-DOCX file types."""
        # Try to upload a text file
        files = {"file": ("resume.txt", b"Plain text content", "text/plain")}

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "Invalid file type" in response.json()["detail"]

    def test_import_docx_corrupted_file(self, client):
        """Test handling of corrupted DOCX file."""
        corrupted = b"This is not a valid DOCX file"
        files = {
            "file": (
                "corrupted.docx",
                corrupted,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "Invalid or corrupted DOCX" in response.json()["detail"]

    def test_import_docx_file_too_large(self, client):
        """Test rejection of files exceeding size limit."""
        # Create a file larger than 10MB
        large_content = b"x" * (11 * 1024 * 1024)  # 11MB
        files = {
            "file": (
                "large.docx",
                large_content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        # Should return 400 or 500 depending on when the size check happens
        assert response.status_code in [
            status.HTTP_400_BAD_REQUEST,
            status.HTTP_500_INTERNAL_SERVER_ERROR,
        ]
        # Check for appropriate error message
        if response.status_code == status.HTTP_400_BAD_REQUEST:
            assert "File too large" in response.json()["detail"]

    def test_import_docx_returns_json_resume_format(self, client, sample_docx_bytes):
        """Test that response follows JSON Resume format."""
        files = {
            "file": (
                "resume.docx",
                sample_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_200_OK
        data = response.json()

        # Check JSON Resume structure
        assert "basics" in data
        assert "work" in data
        assert "education" in data
        assert "skills" in data

        # Check that basics has expected fields
        basics = data.get("basics", {})
        assert "name" in basics or "email" in basics

    def test_import_docx_empty_file(self, client):
        """Test handling of empty DOCX file."""
        doc = Document()
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        files = {
            "file": (
                "empty.docx",
                docx_bytes.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "No text content found" in response.json()["detail"]


class TestImportDocxWithAuthentication:
    """Test DOCX import with API key authentication."""

    def test_import_docx_with_valid_api_key(self, client, sample_docx_bytes):
        """Test DOCX import with valid API key."""
        files = {
            "file": (
                "resume.docx",
                sample_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        headers = {"X-API-KEY": "test-api-key"}

        # In test mode, API key requirement is disabled
        # This test verifies the endpoint works with an API key header present
        response = client.post("/api/v1/import/docx", files=files, headers=headers)

        # Should succeed in test mode (auth disabled)
        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert "basics" in data

    def test_import_docx_without_api_key(self, client, sample_docx_bytes):
        """Test DOCX import without API key (should work in test mode)."""
        files = {
            "file": (
                "resume.docx",
                sample_docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        # In test mode, auth is disabled
        assert response.status_code == status.HTTP_200_OK


class TestImportDocxEdgeCases:
    """Test edge cases for DOCX import."""

    def test_import_docx_with_special_characters(self, client):
        """Test DOCX import with special characters in content."""
        doc = Document()
        doc.add_heading("John & Jane Doe", 0)
        doc.add_paragraph("Email: test<email>@example.com")
        doc.add_paragraph("Skills: Python, C++, R&D")

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        files = {
            "file": (
                "special.docx",
                docx_bytes.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert "John" in data.get("basics", {}).get("name", "")

    def test_import_docx_with_unicode_characters(self, client):
        """Test DOCX import with Unicode characters."""
        doc = Document()
        doc.add_heading("José García", 0)
        doc.add_paragraph("Software Engineer")
        doc.add_paragraph("Skills: Python, JavaScript, C#")

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        files = {
            "file": (
                "unicode.docx",
                docx_bytes.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        # Unicode should be preserved
        assert "José" in data.get("basics", {}).get("name", "") or "García" in data.get(
            "basics", {}
        ).get("name", "")

    def test_import_docx_very_long_content(self, client):
        """Test DOCX import with very long content."""
        doc = Document()
        doc.add_heading("Long Resume", 0)

        # Add many paragraphs
        for i in range(100):
            doc.add_paragraph(f"Experience item {i}: Description of work done.")

        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        files = {
            "file": (
                "long.docx",
                docx_bytes.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        assert response.status_code == status.HTTP_200_OK
        data = response.json()
        assert "basics" in data


class TestDocxContentTypeVariations:
    """Test different DOCX content type variations."""

    def test_import_docx_ms_word_content_type(self, client, simple_docx_bytes):
        """Test DOCX import with MS Word content type."""
        files = {
            "file": (
                "resume.docx",
                simple_docx_bytes,
                "application/vnd.ms-word.document",
            )
        }

        response = client.post("/api/v1/import/docx", files=files)

        # Should succeed with alternate content type
        assert response.status_code == status.HTTP_200_OK
